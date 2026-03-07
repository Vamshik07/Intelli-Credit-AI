from docx import Document

def generate_cam(company, financials, risk, recommendation):

    doc = Document()

    doc.add_heading("Credit Appraisal Memo", level=1)

    doc.add_heading("Company", level=2)
    doc.add_paragraph(company)

    doc.add_heading("Financial Summary", level=2)
    doc.add_paragraph(str(financials))

    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(str(risk))

    doc.add_heading("Recommendation", level=2)
    doc.add_paragraph(str(recommendation))

    doc.save("CAM_Report.docx")