from docx import Document
from pathlib import Path
from datetime import datetime


def _five_cs(financials, risk_level, research_data):
    debt_ratio = financials.get("debt_ratio")
    if debt_ratio is None:
        revenue = financials.get("revenue", 0.0)
        debt_ratio = (financials.get("debt", 0.0) / revenue) if revenue else 1.0

    character = "Strong" if research_data.get("promoter_risk_score", 0.5) < 0.3 else "Watchlist"
    capacity = "Adequate" if financials.get("profit", 0.0) > 0 else "Constrained"
    capital = "Healthy" if debt_ratio < 0.4 else "Leveraged"
    collateral = "To Be Verified"
    conditions = "Favorable" if risk_level == "Low Risk" else "Cautious"

    return {
        "Character": character,
        "Capacity": capacity,
        "Capital": capital,
        "Collateral": collateral,
        "Conditions": conditions,
    }


def generate_cam(company, financials, risk, recommendation, explanation="", research_data=None, analyst_notes=""):
    research_data = research_data or {}
    doc = Document()

    doc.add_heading("Credit Appraisal Memo", level=1)

    doc.add_heading("Company", level=2)
    doc.add_paragraph(company)

    doc.add_heading("Financial Summary", level=2)
    doc.add_paragraph(str(financials))

    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(str(risk))
    doc.add_paragraph(explanation)

    doc.add_heading("Recommendation", level=2)
    doc.add_paragraph(str(recommendation))

    doc.add_heading("Five Cs of Credit", level=2)
    for k, v in _five_cs(financials, risk, research_data).items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Research Signals", level=2)
    doc.add_paragraph(str(research_data))

    if analyst_notes:
        doc.add_heading("Analyst Notes", level=2)
        doc.add_paragraph(analyst_notes)

    reports_dir = Path(__file__).resolve().parents[1] / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = reports_dir / f"CAM_Report_{timestamp}.docx"
    doc.save(output_path)
    return str(output_path)
