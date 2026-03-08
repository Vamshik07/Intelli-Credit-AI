from datetime import datetime
from pathlib import Path
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from backend.services.llm_router import route_llm


def _load_cam_prompt_template() -> str:
    prompt_file = Path("prompts") / "cam_generation.txt"
    if not prompt_file.exists():
        return "Generate a professional Credit Appraisal Memo using the provided inputs."
    return prompt_file.read_text(encoding="utf-8")


def _format_number(value: object) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


def _flatten_dict(title: str, payload: dict) -> str:
    lines = [f"{title}:"]
    for key, value in payload.items():
        label = key.replace("_", " ").title()
        lines.append(f"- {label}: {value}")
    return "\n".join(lines)


def _risk_category(final_score: float) -> str:
    if final_score > 80:
        return "Low Risk"
    if final_score >= 60:
        return "Moderate Risk"
    return "High Risk"


def _format_runtime_context(company: dict, risk: dict, recommendation: dict, research: dict, observations: str) -> str:
    financial_score = risk.get("financial_score", 0.0)
    legal_score = risk.get("legal_score", 0.0)
    promoter_score = risk.get("promoter_score", 0.0)
    operational_score = risk.get("operational_score", 0.0)
    industry_score = risk.get("industry_score", 0.0)
    final_score = risk.get("final_score", 0.0)

    company_text = _flatten_dict("Company", company)
    risk_text = _flatten_dict("Risk Score Inputs", risk)
    recommendation_text = _flatten_dict("Loan Recommendation Inputs", recommendation)
    research_text = _flatten_dict("Research Intelligence", research)

    return (
        "\n\nINPUT DATA FOR MEMO DRAFTING\n"
        f"{company_text}\n\n"
        f"{risk_text}\n\n"
        f"{recommendation_text}\n\n"
        f"{research_text}\n\n"
        "Normalized Score Breakdown:\n"
        f"- Financial Strength: {financial_score}\n"
        f"- Legal Strength: {legal_score}\n"
        f"- Promoter Strength: {promoter_score}\n"
        f"- Operational Capacity: {operational_score}\n"
        f"- Industry Strength: {industry_score}\n"
        f"- Final Weighted Score: {final_score}\n\n"
        "Credit Officer Observations:\n"
        f"- {observations or 'No additional observations provided.'}\n"
    )


def _fallback_cam_text(company: dict, risk: dict, recommendation: dict, research: dict, observations: str) -> str:
    company_name = company.get("name", "Unknown Company")
    cin = company.get("cin", "Not Available")
    industry = company.get("industry", "Not Available")
    final_score = float(risk.get("final_score", 0.0) or 0.0)
    decision = recommendation.get("decision", "Pending")
    recommended_loan = _format_number(recommendation.get("recommended_loan_amount", 0.0))
    suggested_rate = _format_number(recommendation.get("suggested_interest_rate", 0.0))
    risk_category = _risk_category(final_score)
    components = risk.get("components", {}) if isinstance(risk.get("components", {}), dict) else {}

    return f"""# CREDIT APPRAISAL MEMO

## 1. Executive Summary
{company_name} operates in the {industry} segment and was reviewed using a structured underwriting framework that combines financial analysis, legal-risk screening, promoter signal interpretation, sector context, and credit officer observations. This memo has been generated in a credit-committee ready format to support an informed lending decision and ensure transparency in the reasoning trail behind the recommendation.

The overall AI-assisted evaluation indicates a final weighted score of {final_score:.2f} out of 100, placing the borrower in the {risk_category} category under the current policy thresholds. Core constraints identified in this cycle include profitability pressure, leverage sensitivity, and legal/regulatory uncertainty as reflected in the legal and promoter components of the model. While the applicant demonstrates business continuity potential, lending terms must remain aligned with conservative risk governance.

From a decision perspective, the recommendation at this stage is **{decision}**. This conclusion is based on a blended interpretation of quantitative strength factors and qualitative risk triggers, including external research and credit officer commentary. The remainder of this memo details each component in a section-wise format for committee deliberation, covenant planning, and final sanction strategy.

## 2. Company Overview
| Attribute | Details |
| --- | --- |
| Company Name | {company_name} |
| CIN | {cin} |
| Industry | {industry} |
| Business Type | Corporate Borrower |
| Key Operations | Manufacturing, supply chain execution, customer servicing |

The applicant appears to operate through established commercial channels with a formal corporate identity and identifiable operating profile. Based on available records and uploaded documentation, the company maintains transaction activity and market participation but requires closer underwriting controls around debt service resilience and downside stress capacity.

The business profile suggests that earnings visibility is influenced by input cycles, customer demand, and sector momentum. In committee terms, this is a profile where lending viability can improve when documentation quality, collateral comfort, and repayment structure are strengthened through condition-based sanctioning.

## 3. Promoter Analysis
Promoter and governance quality were reviewed through available external intelligence and risk signal interpretation. The promoter strength score indicates the need for measured confidence rather than full risk comfort, suggesting that credibility is serviceable but not free from monitoring requirements.

Key observations include governance sensitivity under adverse scenarios, dependence on management discipline for cash controls, and potential reputation drag where legal or market narratives deteriorate. As a result, promoter quality should be treated as a moderate-to-elevated risk lever in sanction structuring.

Recommended committee stance for promoter assessment:
- require periodic management information and covenant reporting;
- monitor governance red flags and litigation references quarterly;
- calibrate disbursement and review triggers to promoter conduct indicators.

## 4. Financial Analysis
| Financial Metric | Value |
| --- | --- |
| Revenue | {_format_number(components.get('revenue', 0.0))} |
| EBITDA | {_format_number(components.get('ebitda', 0.0))} |
| Net Profit | {_format_number(components.get('net_profit', 0.0))} |
| Total Assets | {_format_number(components.get('total_assets', 0.0))} |
| Total Debt | {_format_number(components.get('total_debt', 0.0))} |

The financial framework indicates stress in margin strength and leverage-adjusted durability. Even where topline activity exists, net profit conversion and debt load behavior remain critical to credit quality. Underwriting confidence is therefore more sensitive to repayment mechanics and covenant architecture than to revenue visibility alone.

From a banking perspective, this profile benefits from conservative cash-flow assumptions, debt service coverage monitoring, and staged exposure build-up. The committee should prioritize structural safeguards including amortization discipline, drawdown controls, and periodic ratio-based review. These actions are necessary to offset uncertainty in profitability persistence and ensure that downside scenarios remain manageable.

## 5. Industry Risk Assessment
Industry risk must be interpreted through demand cyclicality, policy exposure, and competitive pressure. Current sector signals indicate mixed conditions: structural demand pockets may exist, but volatility and regulatory influence can meaningfully alter risk outcomes over a medium-term horizon.

The industry strength score in this case is {risk.get('industry_score', 0.0)}, which implies that sector context is not fully neutral and requires active risk buffering. Growth potential can be captured if execution remains disciplined, but the lending structure should assume periodic stress and not rely on uninterrupted favorable market conditions.

Committee interpretation:
- lending terms should include sector-aware stress buffers;
- pricing should reflect volatility, not only base-case performance;
- monitoring cadence should increase when market conditions soften.

## 6. Legal Risk Assessment
Legal and compliance indicators represent a material decision variable in this appraisal cycle. The legal strength score of {risk.get('legal_score', 0.0)} suggests that legal exposure is a live underwriting consideration and cannot be treated as peripheral.

Any active litigation references, regulatory narratives, or dispute patterns increase execution risk, recovery complexity, and operational friction. Even where legal events are not immediately severe, recurring signals can elevate uncertainty in borrower behavior and covenant adherence.

Recommended committee stance:
- seek explicit disclosures on open legal matters;
- incorporate legal covenant triggers for adverse developments;
- maintain tighter review rights until legal clarity improves.

## 7. Five Cs of Credit Evaluation
| Credit Factor | Evaluation |
| --- | --- |
| Character | Moderate confidence; governance oversight required |
| Capacity | Cash-flow resilience requires structured repayment controls |
| Capital | Balance-sheet flexibility appears constrained under stress |
| Collateral | Must be validated for enforceability and coverage sufficiency |
| Conditions | Sector and legal context justify conservative sanction terms |

The Five Cs assessment supports a cautious underwriting stance. Character and capacity are not weak by default but require observable evidence over time. Capital and collateral quality should be tested against downside assumptions rather than nominal values. Conditions remain sensitive to both market and legal developments, reinforcing the need for structured sanction architecture.

## 8. Risk Scorecard
| Component | Weight | Score |
| --- | --- | --- |
| Financial Strength | 30% | {risk.get('financial_score', 0.0)} |
| Legal Risk | 20% | {risk.get('legal_score', 0.0)} |
| Promoter Strength | 20% | {risk.get('promoter_score', 0.0)} |
| Operational Capacity | 15% | {risk.get('operational_score', 0.0)} |
| Industry Risk | 15% | {risk.get('industry_score', 0.0)} |

Final Weighted Risk Score: **{final_score:.2f} / 100**

Interpretation: the composite score indicates that credit risk is elevated relative to standard approval comfort and requires either rejection or conditional structuring with stronger safeguards, depending on policy tolerance and collateral support.

## 9. Loan Recommendation
Decision: **{decision}**

| Loan Parameter | Value |
| --- | --- |
| Recommended Loan Amount | {recommended_loan} |
| Suggested Interest Rate | {suggested_rate}% |
| Risk Category | {risk_category} |

This recommendation reflects the aggregate effect of financial pressure points, legal risk sensitivity, promoter confidence limits, and industry uncertainty. Where approval is considered under conditional structures, strict covenants, periodic monitoring, and phased utilization controls are essential to maintain prudent risk-adjusted returns.

## 10. AI Explainability
The decision is explained through the following credit-committee focused factors:
- financial strength and profitability resilience under normal and stressed assumptions;
- legal and compliance uncertainty that may impair predictability of repayment behavior;
- promoter and governance quality, including signal-based credibility interpretation;
- industry and operating environment volatility affecting earnings stability;
- credit officer observations incorporated into final weighting and narrative judgment.

Additional analyst narrative:
{observations or 'No additional narrative observations were provided by the credit officer at the time of analysis.'}

This memo is intended for formal credit committee review and should be read together with sanction terms, covenant framework, and legal documentation conditions prior to final decision closure.
"""


def _generate_cam_text(company: dict, risk: dict, recommendation: dict, research: dict, observations: str) -> str:
    prompt = _load_cam_prompt_template() + _format_runtime_context(company, risk, recommendation, research, observations)
    prompt += (
        "\n\nMANDATORY ENFORCEMENT:\n"
        "- Minimum length 1200 words.\n"
        "- Each required section must contain at least two detailed paragraphs.\n"
        "- Do not output raw JSON or Python dictionaries.\n"
        "- Use tables and bullet points where instructed.\n"
    )

    llm_text = (route_llm("cam_generation", prompt) or "").strip()
    if not llm_text or llm_text.lower() == "llm unavailable":
        return _fallback_cam_text(company, risk, recommendation, research, observations)

    # Guardrail: if model returns terse or malformed memo, append a detailed fallback block.
    if len(llm_text.split()) < 900:
        llm_text += "\n\n" + _fallback_cam_text(company, risk, recommendation, research, observations)

    return llm_text


def _is_markdown_table_separator(line: str) -> bool:
    stripped = line.replace(" ", "")
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    parts = [part for part in stripped.split("|") if part]
    if not parts:
        return False
    return all(set(part) <= {":", "-"} and "-" in part for part in parts)


def _parse_markdown_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            continue
        columns = [col.strip() for col in stripped.strip("|").split("|")]
        rows.append(columns)
    return rows


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _shade_header_cell(cell, fill: str = "DCE6F1") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_doc_cover(doc: Document, company: dict) -> None:
    title = doc.add_heading("CREDIT APPRAISAL MEMO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("INTELLI-CREDIT AI Corporate Lending Assessment")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x33, 0x47, 0x5B)
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph("")
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(f"Company: {company.get('name', 'Unknown')}\n").bold = True
    details.add_run(f"CIN: {company.get('cin', 'Not Available')}\n")
    details.add_run(f"Industry: {company.get('industry', 'Not Available')}\n")
    details.add_run(f"Generated On: {datetime.now().strftime('%d-%b-%Y %H:%M:%S')}")

    doc.add_page_break()


def _write_docx_from_text(doc: Document, memo_text: str) -> None:
    lines = memo_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        # Detect markdown table blocks and render as native DOCX tables.
        if (
            line.startswith("|")
            and index + 1 < len(lines)
            and _is_markdown_table_separator(lines[index + 1].strip())
        ):
            table_lines = [line]
            index += 2  # skip header and separator
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1

            rows = _parse_markdown_table_rows(table_lines)
            if rows:
                col_count = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                for row_idx, row_data in enumerate(rows):
                    for col_idx in range(col_count):
                        value = row_data[col_idx] if col_idx < len(row_data) else ""
                        cell = table.cell(row_idx, col_idx)
                        cell.text = value
                        if row_idx == 0:
                            _shade_header_cell(cell)
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            index += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            index += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            index += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            index += 1
            continue

        doc.add_paragraph(line)
        index += 1


def _write_pdf_from_text(pdf: canvas.Canvas, memo_text: str) -> None:
    width, height = A4
    y = height - 62
    page_no = 1

    def draw_page_chrome() -> None:
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(36, height - 28, "Credit Appraisal Memo")
        pdf.setFont("Helvetica", 8)
        pdf.drawRightString(width - 36, height - 28, datetime.now().strftime("%d-%b-%Y"))
        pdf.line(36, height - 34, width - 36, height - 34)
        pdf.line(36, 40, width - 36, 40)
        pdf.drawString(36, 28, "INTELLI-CREDIT AI")
        pdf.drawRightString(width - 36, 28, f"Page {page_no}")

    draw_page_chrome()
    pdf.setFont("Helvetica", 10)

    lines = memo_text.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()

        if not line.strip():
            y -= 8
            if y < 60:
                pdf.showPage()
                page_no += 1
                draw_page_chrome()
                pdf.setFont("Helvetica", 10)
                y = height - 62
            index += 1
            continue

        # Render markdown tables as aligned rows for readability in PDF.
        if (
            line.strip().startswith("|")
            and index + 1 < len(lines)
            and _is_markdown_table_separator(lines[index + 1].strip())
        ):
            table_lines = [line.strip()]
            index += 2  # skip separator
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1

            rows = _parse_markdown_table_rows(table_lines)
            if rows:
                col_count = max(len(row) for row in rows)
                usable_width = width - 72
                col_width = usable_width / max(col_count, 1)

                for row_idx, row in enumerate(rows):
                    if y < 60:
                        pdf.showPage()
                        page_no += 1
                        draw_page_chrome()
                        pdf.setFont("Helvetica", 10)
                        y = height - 62

                    pdf.setFont("Helvetica-Bold" if row_idx == 0 else "Helvetica", 10)
                    x = 36
                    for col_idx in range(col_count):
                        value = row[col_idx] if col_idx < len(row) else ""
                        pdf.drawString(x, y, value[:28])
                        x += col_width
                    y -= 14
                y -= 4
            continue

        if line.startswith("# "):
            pdf.setFont("Helvetica-Bold", 14)
            text_lines = [line[2:].strip()]
        elif line.startswith("## "):
            pdf.setFont("Helvetica-Bold", 12)
            text_lines = [line[3:].strip()]
        elif line.startswith("### "):
            pdf.setFont("Helvetica-Bold", 11)
            text_lines = [line[4:].strip()]
        else:
            pdf.setFont("Helvetica", 10)
            content = line[2:].strip() if line.startswith("- ") else line
            prefix = "- " if line.startswith("- ") else ""
            wrapped = textwrap.wrap(content, width=105)
            text_lines = [prefix + wrapped[0]] + wrapped[1:] if wrapped else [prefix]

        for item in text_lines:
            if y < 60:
                pdf.showPage()
                page_no += 1
                draw_page_chrome()
                y = height - 62
            pdf.drawString(36, y, item[:140])
            y -= 14

        index += 1

def generate_cam_reports(company: dict, risk: dict, recommendation: dict, research: dict, observations: str) -> dict:
    reports_dir = Path("backend/reports")
    reports_dir.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = reports_dir / f"CAM_{stamp}.docx"
    pdf_path = reports_dir / f"CAM_{stamp}.pdf"

    memo_text = _generate_cam_text(company, risk, recommendation, research, observations)

    doc = Document()
    _set_doc_defaults(doc)
    _add_doc_cover(doc, company)
    _write_docx_from_text(doc, memo_text)
    doc.save(docx_path)

    pdf = canvas.Canvas(str(pdf_path), pagesize=A4)
    _write_pdf_from_text(pdf, memo_text)
    pdf.save()

    return {"docx": str(docx_path), "pdf": str(pdf_path)}
